import json
import os
from datetime import datetime
from config import (
    RULESETS, ASSESSMENT_AREAS, RATING_LABELS,
    DRILL_RECOMMENDATIONS, SUBMISSIONS
)
try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib import colors
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


def analyze_timeline_for_assessment(events):
    """Analyze timeline events to generate automatic assessment scores."""
    if not events:
        return {}
    
    # Initialize assessment tracking
    assessment_data = {}
    
    for event in events:
        action = event.get("action", "").lower()
        result = event.get("result", "").lower()
        fighter = event.get("fighter", "")
        positioning = event.get("positioning", "").lower()
        points = event.get("points", 0)
        
        # Skip if not Fighter A (athlete being reviewed)
        if fighter != "Fighter A":
            continue
            
        # Determine success based on points, positioning, and result type
        is_positive = (points > 0 or "excellent" in positioning or "strong" in positioning or 
                      "Advantage" in result or points == 0.5)
        is_negative = (points < 0 or "poor" in positioning or "lost" in positioning or 
                      "failed" in action or "Penalty" in result)
        is_attempt = "Attempt" in result  # Track attempts separately
        
        # Map actions to assessment areas
        area_mappings = {
            "takedown": "Takedowns / Stand-Up",
            "single leg": "Takedowns / Stand-Up", 
            "double leg": "Takedowns / Stand-Up",
            "guard pull": "Guard Pulling",
            "guard pass": "Guard Passing",
            "pass": "Guard Passing",
            "sweep": "Sweeps",
            "submission": "Submissions (Offense)",
            "back take": "Back Takes",
            "mount": "Mount Offense", 
            "side control": "Side Control Offense",
            "escape": determine_escape_type(action),
            "grip": "Grip Fighting",
            "transition": "Transitions"
        }
        
        # Find matching assessment area
        for keyword, area in area_mappings.items():
            if keyword in action:
                if area and area in ASSESSMENT_AREAS:
                    if area not in assessment_data:
                        assessment_data[area] = {"attempts": 0, "successes": 0, "failures": 0, "advantages": 0}
                    
                    assessment_data[area]["attempts"] += 1
                    if is_positive and not is_attempt:
                        assessment_data[area]["successes"] += 1
                    elif is_attempt and points > 0:  # Attempt that earned advantage
                        assessment_data[area]["advantages"] += 1
                    elif is_negative:
                        assessment_data[area]["failures"] += 1
                break
    
    # Generate scores (1-5) based on success rates and attempt counts
    final_assessments = {}
    for area, data in assessment_data.items():
        attempts = data["attempts"]
        successes = data["successes"] 
        failures = data["failures"]
        advantages = data["advantages"]
        neutrals = attempts - successes - failures - advantages
        
        if attempts == 0:
            continue
            
        # Calculate score considering full successes, advantages, and failures
        success_rate = successes / attempts if attempts > 0 else 0
        advantage_rate = advantages / attempts if attempts > 0 else 0
        failure_rate = failures / attempts if attempts > 0 else 0
        effective_success_rate = success_rate + (advantage_rate * 0.5)  # Weight advantages as half success
        
        if effective_success_rate >= 0.8 or success_rate >= 0.7:
            score = 5  # Elite
        elif effective_success_rate >= 0.6 or success_rate >= 0.5:
            score = 4  # Strong  
        elif failure_rate <= 0.2 and effective_success_rate >= 0.3:
            score = 3  # Competent
        elif failure_rate <= 0.5:
            score = 2  # Needs Work
        else:
            score = 1  # Significant Weakness
            
        final_assessments[area] = score
    
    return final_assessments


def generate_ai_summary_report(review_data):
    """Generate AI-powered summary report with training recommendations."""
    timeline = review_data.get("timeline", [])
    assessments = review_data.get("assessments", {})
    analysis = review_data.get("analysis", {})
    tactical_tags = review_data.get("tactical_tags", [])
    scores = review_data.get("scores", {})
    
    report = []
    
    # Performance Overview
    report.append("# 🤖 AI SUMMARY REPORT")
    report.append("")
    
    final_score = f"A: {scores.get('fighter_a', 0)} | B: {scores.get('fighter_b', 0)}"
    fighter_a_score = scores.get('fighter_a', 0)
    fighter_b_score = scores.get('fighter_b', 0)
    
    if fighter_a_score > fighter_b_score:
        outcome = "Won by points - Strong performance with room for optimization"
    elif fighter_a_score < fighter_b_score:
        outcome = "Lost by points - Key areas identified for immediate improvement"
    else:
        outcome = "Draw/Close match - Fine margins determine outcome"
    
    report.append(f"**Match Outcome:** {outcome}")
    report.append(f"**Final Score:** {final_score}")
    report.append("")
    
    # Key Insights
    report.append("## 🔍 Key Performance Insights")
    
    # Analyze assessment distribution
    if assessments:
        high_scores = [area for area, data in assessments.items() if data.get('rating', 0) >= 4]
        low_scores = [area for area, data in assessments.items() if data.get('rating', 0) <= 2]
        
        if high_scores:
            report.append(f"**Dominant Areas:** {', '.join(high_scores[:3])}")
        if low_scores:
            report.append(f"**Critical Gaps:** {', '.join(low_scores[:3])}")
    
    # Analyze timeline patterns
    if timeline:
        fighter_a_events = [e for e in timeline if e.get('fighter') == 'Fighter A']
        total_points = sum(e.get('points', 0) for e in fighter_a_events)
        positive_events = [e for e in fighter_a_events if e.get('points', 0) > 0]
        negative_events = [e for e in fighter_a_events if e.get('points', 0) < 0]
        missed_opps = [e for e in fighter_a_events if e.get('missed_opportunity') and e.get('missed_opportunity') != 'None']
        
        if positive_events:
            report.append(f"**Scoring Moments:** {len(positive_events)} successful scoring actions")
        if negative_events:
            report.append(f"**Penalty Issues:** {len(negative_events)} penalties/negative points")
        if missed_opps:
            report.append(f"**Missed Opportunities:** {len(missed_opps)} identified chances for improvement")
    
    report.append("")
    
    # Missed Opportunities Analysis
    if timeline:
        missed_opps = [e.get('missed_opportunity') for e in timeline 
                      if e.get('fighter') == 'Fighter A' and e.get('missed_opportunity') and e.get('missed_opportunity') != 'None']
        
        if missed_opps:
            report.append("## 🎯 Key Missed Opportunities")
            # Count most common missed opportunities
            from collections import Counter
            opp_counts = Counter(missed_opps)
            top_missed = opp_counts.most_common(3)
            
            for opp, count in top_missed:
                report.append(f"**{opp}** (occurred {count} time{'s' if count != 1 else ''})")
                # Add specific recommendations based on missed opportunity type
                if "scored points" in opp.lower():
                    report.append("   • Focus on finishing sequences with authority")
                    report.append("   • Practice high-percentage scoring positions")
                elif "improved position" in opp.lower():
                    report.append("   • Work on position advancement chains")
                    report.append("   • Study transition timing from timeline moments")
                elif "submission" in opp.lower():
                    report.append("   • Drill submission setups with more precision")
                    report.append("   • Practice maintaining control during finish attempts")
                elif "defended" in opp.lower():
                    report.append("   • Improve defensive awareness and reactions")
                    report.append("   • Practice early recognition of dangerous positions")
            report.append("")
    
    
    if assessments:
        # Get lowest scoring areas for priority focus
        priority_areas = sorted(
            [(area, data.get('rating', 0)) for area, data in assessments.items()],
            key=lambda x: x[1]
        )[:3]
        
        for i, (area, rating) in enumerate(priority_areas, 1):
            urgency = "HIGH" if rating <= 2 else "MEDIUM" if rating == 3 else "LOW"
            report.append(f"**{i}. {area}** (Priority: {urgency})")
            
            # Add specific recommendations based on area
            recommendations = get_training_recommendations(area, rating, timeline)
            for rec in recommendations:
                report.append(f"   • {rec}")
            report.append("")
    
    # Tactical Patterns
    if tactical_tags:
        report.append("## 📊 Tactical Pattern Analysis")
        
        # Group tactical tags by category
        positive_tags = [tag for tag in tactical_tags if any(word in tag.lower() 
                        for word in ['strong', 'excellent', 'good', 'superior', 'effective'])]
        negative_tags = [tag for tag in tactical_tags if any(word in tag.lower() 
                        for word in ['poor', 'weak', 'failed', 'lost', 'late', 'predictable'])]
        
        if positive_tags:
            report.append("**Consistent Strengths:**")
            for tag in positive_tags[:3]:
                report.append(f"   • {tag}")
        
        if negative_tags:
            report.append("**Recurring Issues:**")
            for tag in negative_tags[:3]:
                report.append(f"   • {tag}")
        report.append("")
    
    # Next Match Game Plan
    report.append("## 🎮 Next Match Game Plan")
    report.append("**Primary Strategy:** Focus on proven strengths while shoring up critical gaps")
    
    if assessments:
        strong_areas = [area for area, data in assessments.items() if data.get('rating', 0) >= 4]
        if strong_areas:
            report.append(f"**Lean Into:** {strong_areas[0]} - Your most reliable weapon")
        
        weak_areas = [area for area, data in assessments.items() if data.get('rating', 0) <= 2]
        if weak_areas:
            report.append(f"**Avoid Early:** {weak_areas[0]} - Work this in training first")
    
    report.append("")
    report.append("**Pre-Match Preparation:**")
    report.append("   • Drill top 2 priority areas for 15 minutes before competing")
    report.append("   • Visualization: Successfully executing your strongest techniques")
    report.append("   • Game plan: Lead with strengths, use improved areas situationally")
    report.append("")
    
    # Training Schedule Recommendation
    report.append("## 📅 Recommended Training Split")
    if assessments:
        weak_count = len([a for a, d in assessments.items() if d.get('rating', 0) <= 2])
        if weak_count >= 3:
            report.append("**Intensive Phase:** 70% fundamentals, 30% advanced techniques")
        elif weak_count >= 1:
            report.append("**Balanced Phase:** 50% skill gaps, 50% strength development")
        else:
            report.append("**Competition Phase:** 30% weaknesses, 70% strength refinement")
    
    report.append("**Weekly Focus:** 3x positional drilling, 2x live rolling, 1x specific techniques")
    report.append("")
    
    return "\n".join(report)


def get_training_recommendations(area, rating, timeline):
    """Get specific training recommendations for an assessment area."""
    recommendations = {
        "Takedowns / Stand-Up": [
            "2x weekly takedown drilling with specific timing focus",
            "Study opponent's stance and grip preferences on video",
            "Chain drilling: setup > shot > finish > recovery"
        ],
        "Guard Passing": [
            "Headquarters position maintenance drill (10 min daily)",
            "Pressure passing vs speed passing situation awareness",
            "Study guard player's retention patterns from timeline"
        ],
        "Guard Retention": [
            "Hip escape chain drilling with decreasing space",
            "Frame battle conditioning (3x5 min rounds)",
            "Re-guard recovery from bad positions"
        ],
        "Submissions (Offense)": [
            "Submission chain flow: primary > backup > position",
            "Setup drilling without resistance for timing",
            "Study missed opportunities from timeline events"
        ],
        "Submission Defense": [
            "Escape drilling from bad positions under pressure",
            "Hand fighting and posture maintenance basics",
            "Survival position training in worst-case scenarios"
        ]
    }
    
    # Add urgency modifiers based on rating
    base_recs = recommendations.get(area, [
        "Focus on drilling fundamentals with resistance",
        "Positional sparring from disadvantageous positions",
        "Video study of high-level athletes in this area"
    ])
    
    if rating <= 2:
        base_recs.append("DAILY practice recommended - highest priority")
    
    return base_recs[:3]  # Limit to 3 recommendations


def determine_escape_type(action):
    """Determine specific escape type from action text."""
    escape_types = {
        "mount": "Mount Escapes",
        "side control": "Side Control Escapes", 
        "back": "Back Escapes"
    }
    
    for keyword, escape_type in escape_types.items():
        if keyword in action:
            return escape_type
    return None


def calculate_score(events, ruleset_key):
    """Calculate running score from timeline events based on explicit points."""
    fighter_a_score = 0
    fighter_b_score = 0
    
    for event in events:
        fighter = event.get("fighter", "")
        points = event.get("points", 0)  # Use explicit points from event
        
        if fighter == "Fighter A":
            fighter_a_score += points
        elif fighter == "Fighter B":
            fighter_b_score += points
    
    return {
        "fighter_a": round(fighter_a_score, 1),
        "fighter_b": round(fighter_b_score, 1)
    }


def generate_strengths_weaknesses(assessments):
    """Auto-generate strengths and weaknesses from assessment ratings."""
    strengths = []
    weaknesses = []
    
    for area, rating in assessments.items():
        if rating >= 4:
            strengths.append(f"{area} ({RATING_LABELS.get(rating, '')})")
        elif rating <= 2:
            weaknesses.append(f"{area} ({RATING_LABELS.get(rating, '')})")
    
    return strengths, weaknesses


def get_drill_recommendations(assessments, threshold=2):
    """Get drill recommendations for areas rated at or below the threshold."""
    recommendations = {}
    
    for area, rating in assessments.items():
        if rating <= threshold and area in DRILL_RECOMMENDATIONS:
            recommendations[area] = {
                "rating": rating,
                "rating_label": RATING_LABELS.get(rating, ""),
                "drills": DRILL_RECOMMENDATIONS[area]
            }
    
    return recommendations


def get_next_review_id():
    """Generate the next sequential review ID starting at 0001."""
    # Check both data/reviews and reviews directories for existing files
    review_dirs = ["reviews", os.path.join("data", "reviews")]
    
    # Also check user-specific review directories
    if os.path.exists(os.path.join("data", "users")):
        for user_folder in os.listdir(os.path.join("data", "users")):
            user_review_dir = os.path.join("data", "users", user_folder, "reviews")
            if os.path.exists(user_review_dir):
                review_dirs.append(user_review_dir)
    
    highest_id = 0
    
    for review_dir in review_dirs:
        if not os.path.exists(review_dir):
            continue
            
        for filename in os.listdir(review_dir):
            if filename.endswith('.json'):
                # Extract ID from different formats
                if filename.startswith('REV-') and len(filename) >= 8:
                    # New format: REV-0001.json
                    id_part = filename[4:8]  # Extract 4 digit number
                    if id_part.isdigit():
                        highest_id = max(highest_id, int(id_part))
                elif filename.startswith('BJJ-'):
                    # Legacy format: BJJ-20260411-004006.json
                    # Extract timestamp and convert to sequential ID for compatibility
                    try:
                        # Parse existing JSON to check if it has a sequential review_id
                        filepath = os.path.join(review_dir, filename)
                        with open(filepath, 'r') as f:
                            data = json.load(f)
                            review_id = data.get('metadata', {}).get('review_id', '')
                            if review_id.startswith('REV-') and len(review_id) >= 8:
                                id_part = review_id[4:8]
                                if id_part.isdigit():
                                    highest_id = max(highest_id, int(id_part))
                    except:
                        pass
    
    next_id = highest_id + 1
    return f"REV-{next_id:04d}"


def build_review_data(match_info, events, assessments, tactical_tags, notes, custom_strengths="", custom_improvements="", existing_review_id=None):
    """Build the complete review data structure."""
    scores = calculate_score(events, match_info.get("ruleset", "IBJJF"))
    strengths, weaknesses = generate_strengths_weaknesses(assessments)
    drill_recs = get_drill_recommendations(assessments)
    
    # Use existing review ID if editing, otherwise generate new one
    review_id = existing_review_id if existing_review_id else get_next_review_id()
    
    review = {
        "metadata": {
            "review_id": review_id,
            "reviewed_at": datetime.now().isoformat(),
            "reviewer": "BJJ Mat IQ",
            "platform_version": "Phase 2.0"
        },
        "match_info": match_info,
        "timeline": events,
        "scores": scores,
        "assessments": {
            area: {
                "rating": rating,
                "label": RATING_LABELS.get(rating, "")
            }
            for area, rating in assessments.items()
        },
        "tactical_tags": tactical_tags,
        "analysis": {
            "strengths": strengths,
            "weaknesses": weaknesses,
            "drill_recommendations": drill_recs,
            "custom_strengths": custom_strengths,
            "custom_improvements": custom_improvements,
            "ai_summary": ""  # Will be populated when generated
        },
        "coach_notes": notes
    }
    
    return review


def export_json(review_data, output_dir="reviews"):
    """Export review data as JSON file."""
    os.makedirs(output_dir, exist_ok=True)
    review_id = review_data["metadata"]["review_id"]
    filepath = os.path.join(output_dir, f"{review_id}.json")
    
    with open(filepath, "w") as f:
        json.dump(review_data, f, indent=2)
    
    return filepath


def export_markdown(review_data):
    """Export review data as Markdown report string."""
    info = review_data["match_info"]
    scores = review_data["scores"]
    analysis = review_data["analysis"]
    timeline = review_data["timeline"]
    assessments = review_data["assessments"]
    meta = review_data["metadata"]
    
    md = []
    md.append(f"# BJJ Match Review Report")
    md.append(f"**Review ID:** {meta['review_id']}")
    md.append(f"**Date:** {meta['reviewed_at'][:10]}")
    md.append(f"**Reviewer:** {meta['reviewer']}")
    md.append("")
    
    md.append("## Match Information")
    md.append(f"- **Fighter A:** {info.get('fighter_a', 'N/A')}")
    if info.get('team_a'):
        md.append(f"  - **Team:** {info.get('team_a', 'N/A')}")
    md.append(f"- **Fighter B:** {info.get('fighter_b', 'N/A')}")
    if info.get('team_b'):
        md.append(f"  - **Team:** {info.get('team_b', 'N/A')}")
    md.append(f"- **Belt Level:** {info.get('belt', 'N/A')}")
    md.append(f"- **Weight Class:** {info.get('weight_class', 'N/A')}")
    md.append(f"- **Gi / No-Gi:** {info.get('gi_nogi', 'N/A')}")
    md.append(f"- **Ruleset:** {info.get('ruleset', 'N/A')}")
    md.append(f"- **Event:** {info.get('event', 'N/A')}")
    md.append(f"- **Video URL:** {info.get('video_url', 'N/A')}")
    md.append(f"- **Result:** {info.get('result', 'N/A')}")
    md.append("")
    
    md.append("## Final Score")
    md.append(f"- **{info.get('fighter_a', 'Fighter A')}:** {scores['fighter_a']} points")
    md.append(f"- **{info.get('fighter_b', 'Fighter B')}:** {scores['fighter_b']} points")
    md.append("")
    
    md.append("## Key Moments Timeline")
    md.append("| Time | Fighter | Position | Action | Result | Points | Notes | Why? | Positioning | Missed Opportunity |")
    md.append("|------|---------|----------|--------|--------|--------|-------|------|-------------|-------------------|")
    for event in timeline:
        time_str = event.get("timestamp", "")
        fighter = event.get("fighter", "")
        position = event.get("position", "")
        action = event.get("action", "")
        result = event.get("result", "")
        points = event.get("points", 0)
        note = event.get("notes", "")
        why = event.get("why", "")
        positioning = event.get("positioning", "")
        missed_opp = event.get("missed_opportunity", "")
        md.append(f"| {time_str} | {fighter} | {position} | {action} | {result} | {points} | {note} | {why} | {positioning} | {missed_opp} |")
    md.append("")
    
    md.append("## Assessment")
    md.append("| Area | Rating | Level |")
    md.append("|------|--------|-------|")
    for area, data in assessments.items():
        bar = "★" * data["rating"] + "☆" * (5 - data["rating"])
        md.append(f"| {area} | {bar} {data['rating']}/5 | {data['label']} |")
    md.append("")
    
    md.append("## Strengths")
    for s in analysis["strengths"]:
        md.append(f"- ✅ {s}")
    if analysis.get("custom_strengths"):
        md.append("")
        md.append("### Additional Strengths")
        md.append(analysis["custom_strengths"])
    md.append("")
    
    md.append("## Areas for Improvement")
    for w in analysis["weaknesses"]:
        md.append(f"- ⚠️ {w}")
    if analysis.get("custom_improvements"):
        md.append("")
        md.append("### Additional Areas for Improvement")
        md.append(analysis["custom_improvements"])
    md.append("")
    
    if analysis["drill_recommendations"]:
        md.append("## Recommended Drills")
        for area, rec_data in analysis["drill_recommendations"].items():
            md.append(f"### {area} (rated {rec_data['rating']}/5 - {rec_data['rating_label']})")
            for drill in rec_data["drills"]:
                md.append(f"- 🔄 {drill}")
            md.append("")
    
    if review_data.get("tactical_tags"):
        md.append("## Tactical Observations")
        for tag in review_data["tactical_tags"]:
            md.append(f"- 🏷️ {tag}")
        md.append("")
    
    if review_data.get("coach_notes"):
        md.append("## Coach Notes")
        md.append(review_data["coach_notes"])
        md.append("")
    
    # Add AI Summary Report
    ai_summary = generate_ai_summary_report(review_data)
    md.append(ai_summary)
    md.append("")
    
    md.append("---")
    md.append("*Generated by BJJ Mat IQ - Phase 1 Review Tool*")
    
    return "\n".join(md)


def export_pdf(review_data, output_dir=None):
    """Export review data as PDF document, return bytes for download."""
    if not REPORTLAB_AVAILABLE:
        return None
    
    from io import BytesIO
    
    info = review_data["match_info"]
    scores = review_data["scores"]
    analysis = review_data["analysis"]
    timeline = review_data["timeline"]
    assessments = review_data["assessments"]
    meta = review_data["metadata"]
    
    # Create PDF in memory
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4)
    story = []
    styles = getSampleStyleSheet()
    
    # Custom styles
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=18,
        spaceAfter=20,
        textColor=colors.darkblue
    )
    
    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontSize=14,
        spaceBefore=15,
        spaceAfter=10,
        textColor=colors.darkblue
    )
    
    # Title
    story.append(Paragraph("BJJ Assessment Analysis Report", title_style))
    story.append(Spacer(1, 12))
    
    # Match Information
    story.append(Paragraph("Match Information", heading_style))
    match_info_data = [
        ["Review ID:", meta['review_id']],
        ["Date:", meta['reviewed_at'][:10]],
        ["Fighter A:", info.get('fighter_a', 'N/A')],
        ["Team A:", info.get('team_a', 'N/A') if info.get('team_a') else 'N/A'],
        ["Fighter B:", info.get('fighter_b', 'N/A')],
        ["Team B:", info.get('team_b', 'N/A') if info.get('team_b') else 'N/A'],
        ["Belt Level:", info.get('belt', 'N/A')],
        ["Weight Class:", info.get('weight_class', 'N/A')],
        ["Gi / No-Gi:", info.get('gi_nogi', 'N/A')],
        ["Ruleset:", info.get('ruleset', 'N/A')],
        ["Event:", info.get('event', 'N/A') if info.get('event') else 'Assessment'],
        ["Result:", info.get('result', 'N/A')]
    ]
    
    match_table = Table(match_info_data, colWidths=[2*inch, 4*inch])
    match_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (0, -1), colors.lightgrey),
        ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ('GRID', (0, 0), (-1, -1), 1, colors.black)
    ]))
    story.append(match_table)
    story.append(Spacer(1, 15))
    
    # Final Score
    story.append(Paragraph("Final Score", heading_style))
    score_text = f"<b>{info.get('fighter_a', 'Fighter A')}:</b> {scores['fighter_a']} points<br/>"
    score_text += f"<b>{info.get('fighter_b', 'Fighter B')}:</b> {scores['fighter_b']} points"
    story.append(Paragraph(score_text, styles['Normal']))
    story.append(Spacer(1, 15))
    
    # Assessment
    story.append(Paragraph("Performance Assessment", heading_style))
    assessment_data = [["Skill Area", "Rating", "Level"]]
    for area, data in assessments.items():
        stars = "★" * data["rating"] + "☆" * (5 - data["rating"])
        assessment_data.append([area, f"{stars} ({data['rating']}/5)", data['label']])
    
    assessment_table = Table(assessment_data, colWidths=[3*inch, 2*inch, 1.5*inch])
    assessment_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.darkblue),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 0), (-1, -1), 9),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.lightgrey])
    ]))
    story.append(assessment_table)
    story.append(Spacer(1, 15))
    
    # Strengths
    story.append(Paragraph("Strengths", heading_style))
    for strength in analysis.get("strengths", []):
        story.append(Paragraph(f"• {strength}", styles['Normal']))
    if analysis.get("custom_strengths"):
        story.append(Paragraph("Additional Strengths:", styles['Heading3']))
        story.append(Paragraph(analysis["custom_strengths"], styles['Normal']))
    story.append(Spacer(1, 10))
    
    # Areas for Improvement
    story.append(Paragraph("Areas for Improvement", heading_style))
    for weakness in analysis.get("weaknesses", []):
        story.append(Paragraph(f"• {weakness}", styles['Normal']))
    if analysis.get("custom_improvements"):
        story.append(Paragraph("Additional Areas for Improvement:", styles['Heading3']))
        story.append(Paragraph(analysis["custom_improvements"], styles['Normal']))
    story.append(Spacer(1, 15))
    
    # Drill Recommendations
    if analysis.get("drill_recommendations"):
        story.append(Paragraph("Recommended Training Drills", heading_style))
        for area, rec_data in analysis["drill_recommendations"].items():
            story.append(Paragraph(f"{area} (Rating: {rec_data['rating']}/5 - {rec_data['rating_label']})", styles['Heading3']))
            for drill in rec_data["drills"]:
                story.append(Paragraph(f"• {drill}", styles['Normal']))
            story.append(Spacer(1, 8))
    
    # Timeline (complete)
    if timeline:
        story.append(Paragraph("Key Match Moments", heading_style))
        timeline_data = [["Time", "Action", "Result", "Points", "Notes"]]
        for event in timeline:
            notes = event.get("notes", "")
            timeline_data.append([
                event.get("timestamp", ""),
                event.get("action", ""),
                event.get("result", ""),
                str(event.get("points", 0)),
                notes
            ])
        
        timeline_table = Table(timeline_data, colWidths=[0.8*inch, 2*inch, 1.5*inch, 0.8*inch, 1.9*inch])
        timeline_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.darkblue),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('ALIGN', (4, 0), (4, -1), 'LEFT'),  # Notes column left-aligned
            ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
            ('FONTSIZE', (0, 0), (-1, -1), 8),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.lightgrey])
        ]))
        story.append(timeline_table)
    
    story.append(Spacer(1, 20))
    story.append(Paragraph("Generated by BJJ Mat IQ - Training Analysis System", styles['Italic']))
    
    # Build PDF and return bytes
    doc.build(story)
    buffer.seek(0)
    return buffer.getvalue()


def export_word(review_data, output_dir=None):
    """Export review data as Word document, return bytes for download."""
    if not DOCX_AVAILABLE:
        return None
    
    from io import BytesIO
    
    info = review_data["match_info"]
    scores = review_data["scores"]
    analysis = review_data["analysis"]
    timeline = review_data["timeline"]
    assessments = review_data["assessments"]
    meta = review_data["metadata"]
    
    # Create document
    doc = Document()
    
    # Title
    title = doc.add_heading('BJJ Assessment Analysis Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Match Information
    doc.add_heading('Match Information', level=1)
    match_table = doc.add_table(rows=12, cols=2)
    match_table.style = 'Table Grid'
    
    match_info_data = [
        ("Review ID:", meta['review_id']),
        ("Date:", meta['reviewed_at'][:10]),
        ("Fighter A:", info.get('fighter_a', 'N/A')),
        ("Team A:", info.get('team_a', 'N/A') if info.get('team_a') else 'N/A'),
        ("Fighter B:", info.get('fighter_b', 'N/A')),
        ("Team B:", info.get('team_b', 'N/A') if info.get('team_b') else 'N/A'),
        ("Belt Level:", info.get('belt', 'N/A')),
        ("Weight Class:", info.get('weight_class', 'N/A')),
        ("Gi / No-Gi:", info.get('gi_nogi', 'N/A')),
        ("Ruleset:", info.get('ruleset', 'N/A')),
        ("Event:", info.get('event', 'N/A') if info.get('event') else 'Assessment'),
        ("Result:", info.get('result', 'N/A'))
    ]
    
    for i, (label, value) in enumerate(match_info_data):
        match_table.cell(i, 0).text = label
        match_table.cell(i, 1).text = str(value)
        # Bold the labels
        match_table.cell(i, 0).paragraphs[0].runs[0].font.bold = True
    
    # Final Score
    doc.add_heading('Final Score', level=1)
    doc.add_paragraph(f"{info.get('fighter_a', 'Fighter A')}: {scores['fighter_a']} points")
    doc.add_paragraph(f"{info.get('fighter_b', 'Fighter B')}: {scores['fighter_b']} points")
    
    # Performance Assessment
    doc.add_heading('Performance Assessment', level=1)
    assessment_table = doc.add_table(rows=len(assessments)+1, cols=3)
    assessment_table.style = 'Table Grid'
    
    # Headers
    hdr_cells = assessment_table.rows[0].cells
    hdr_cells[0].text = 'Skill Area'
    hdr_cells[1].text = 'Rating'
    hdr_cells[2].text = 'Level'
    
    for hdr_cell in hdr_cells:
        for paragraph in hdr_cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    # Data
    for i, (area, data) in enumerate(assessments.items(), 1):
        row_cells = assessment_table.rows[i].cells
        stars = "★" * data["rating"] + "☆" * (5 - data["rating"])
        row_cells[0].text = area
        row_cells[1].text = f"{stars} ({data['rating']}/5)"
        row_cells[2].text = data['label']
        
        # Color code based on rating
        if data["rating"] >= 4:
            # Green for strong
            pass
        elif data["rating"] >= 3:
            # Yellow for competent
            pass
        else:
            # Red for needs work
            pass
    
    # Strengths
    doc.add_heading('Strengths', level=1)
    for strength in analysis.get("strengths", []):
        p = doc.add_paragraph()
        p.add_run("• ").font.bold = True
        p.add_run(strength)
    
    if analysis.get("custom_strengths"):
        doc.add_heading('Additional Strengths', level=2)
        doc.add_paragraph(analysis["custom_strengths"])
    
    # Areas for Improvement
    doc.add_heading('Areas for Improvement', level=1)
    for weakness in analysis.get("weaknesses", []):
        p = doc.add_paragraph()
        p.add_run("• ").font.bold = True
        p.add_run(weakness)
    
    if analysis.get("custom_improvements"):
        doc.add_heading('Additional Areas for Improvement', level=2)
        doc.add_paragraph(analysis["custom_improvements"])
    
    # Drill Recommendations
    if analysis.get("drill_recommendations"):
        doc.add_heading('Recommended Training Drills', level=1)
        for area, rec_data in analysis["drill_recommendations"].items():
            doc.add_heading(f"{area} (Rating: {rec_data['rating']}/5 - {rec_data['rating_label']})", level=2)
            for drill in rec_data["drills"]:
                p = doc.add_paragraph()
                p.add_run("• ").font.bold = True
                p.add_run(drill)
    
    # Timeline Summary
    if timeline:
        doc.add_heading('Key Match Moments', level=1)
        doc.add_paragraph("Timeline of significant events during the match:")
        
        for i, event in enumerate(timeline, 1):
            p = doc.add_paragraph()
            p.add_run(f"{i}. ").font.bold = True
            p.add_run(f"Time {event.get('timestamp', 'N/A')} - ")
            p.add_run(f"{event.get('action', '')}: ").font.bold = True
            p.add_run(f"{event.get('result', '')} ({event.get('points', 0)} points)")
            
            if event.get('notes'):
                p.add_run(f"\n   Note: {event.get('notes', '')}")
    
    # Footer
    doc.add_page_break()
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_p.add_run("Generated by BJJ Mat IQ - Training Analysis System").font.italic = True
    
    # Save to bytes and return
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()